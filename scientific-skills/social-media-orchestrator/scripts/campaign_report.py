#!/usr/bin/env python3
"""
Campaign Report Generator

Aggregates campaign plan data and analytics into a comprehensive
performance report. Supports Markdown and DOCX output formats.

Usage:
    python campaign_report.py \
        --campaign-plan ./campaign_output/campaign_plan.json \
        --analytics ./campaign_output/analytics/report_30d.json \
        --output ./campaign_output/campaign_report.md \
        --format md
"""

import argparse
import json
import os
import sys
from datetime import datetime


# ---------------------------------------------------------------------------
# Report Sections
# ---------------------------------------------------------------------------

def build_header(plan: dict, analytics: dict) -> str:
    """Build report header with campaign metadata."""
    campaign_id = plan.get("campaign_id", "unknown")
    campaign_type = plan.get("campaign_type", "unknown").replace("_", " ").title()
    brand_url = plan.get("brand_url", "N/A")
    start = plan.get("timeline", {}).get("start_date", "N/A")
    end = plan.get("timeline", {}).get("end_date", "N/A")
    duration = plan.get("timeline", {}).get("duration_days", "N/A")
    platforms = ", ".join(plan.get("platforms", []))

    return f"""# Campaign Performance Report

**Campaign ID**: {campaign_id}
**Campaign Type**: {campaign_type}
**Brand**: {brand_url}
**Platforms**: {platforms}
**Duration**: {start} to {end} ({duration} days)
**Report Generated**: {datetime.now().strftime("%Y-%m-%d %H:%M")}

---
"""


def build_executive_summary(plan: dict, analytics: dict) -> str:
    """Build executive summary section."""
    total_planned = plan.get("content_plan", {}).get("total_content_pieces", 0)
    goals = plan.get("goals", {})

    summary_data = analytics.get("summary", {})
    total_published = summary_data.get("total_published", 0)
    total_reach = summary_data.get("total_reach", 0)
    total_engagement = summary_data.get("total_engagement", 0)
    avg_engagement_rate = summary_data.get("avg_engagement_rate", 0)
    follower_growth = summary_data.get("follower_growth", 0)

    section = "## Executive Summary\n\n"

    section += "| Metric | Target | Actual | Status |\n"
    section += "|--------|--------|--------|--------|\n"

    if "engagement_rate" in goals:
        target_er = goals["engagement_rate"]
        status = "On Track" if avg_engagement_rate >= float(target_er.replace("%", "")) else "Below Target"
        section += f"| Engagement Rate | {target_er} | {avg_engagement_rate:.1f}% | {status} |\n"

    if "followers" in goals:
        target_f = goals["followers"]
        status = "On Track" if follower_growth >= abs(int(target_f.replace("+", ""))) else "Below Target"
        section += f"| Follower Growth | {target_f} | +{follower_growth} | {status} |\n"

    section += f"| Content Published | {total_planned} planned | {total_published} published | {'Complete' if total_published >= total_planned else 'In Progress'} |\n"
    section += f"| Total Reach | -- | {total_reach:,} | -- |\n"
    section += f"| Total Engagement | -- | {total_engagement:,} | -- |\n"

    section += "\n"
    return section


def build_content_overview(plan: dict, analytics: dict) -> str:
    """Build content production overview section."""
    content_plan = plan.get("content_plan", {})
    mix = content_plan.get("content_mix", {})
    slots = content_plan.get("slots_per_platform", {})

    section = "## Content Production Overview\n\n"

    section += "### Content Mix\n\n"
    section += "| Category | Planned | Description |\n"
    section += "|----------|---------|-------------|\n"
    section += f"| Value | {mix.get('value', 0)} posts | Educational, entertaining, inspiring |\n"
    section += f"| Curated | {mix.get('curated', 0)} posts | Industry news, UGC, shared |\n"
    section += f"| Promotional | {mix.get('promotional', 0)} posts | Direct sales, offers |\n"
    section += f"| **Total** | **{sum(mix.values())}** | |\n\n"

    section += "### Content by Platform\n\n"
    section += "| Platform | Format | Planned Count |\n"
    section += "|----------|--------|---------------|\n"
    for platform, formats in slots.items():
        for fmt, count in formats.items():
            section += f"| {platform.title()} | {fmt} | {count} |\n"
    section += "\n"

    return section


def build_phase_review(plan: dict, analytics: dict) -> str:
    """Build phase-by-phase review section."""
    phases = plan.get("phases", [])
    phase_analytics = analytics.get("phases", {})

    section = "## Phase-by-Phase Review\n\n"

    for phase in phases:
        name = phase["name"].replace("_", " ").title()
        section += f"### {name} Phase\n\n"
        section += f"- **Dates**: {phase['start_date']} to {phase['end_date']} ({phase['duration_days']} days)\n"
        section += f"- **Focus**: {phase['content_focus']}\n"

        pa = phase_analytics.get(phase["name"], {})
        if pa:
            section += f"- **Posts Published**: {pa.get('posts_published', 'N/A')}\n"
            section += f"- **Avg Engagement Rate**: {pa.get('avg_engagement_rate', 'N/A')}\n"
            section += f"- **Top Post**: {pa.get('top_post', 'N/A')}\n"
        else:
            section += "- *Analytics data not yet available for this phase.*\n"
        section += "\n"

    return section


def build_platform_breakdown(analytics: dict) -> str:
    """Build per-platform performance breakdown."""
    platforms = analytics.get("platforms", {})
    if not platforms:
        return "## Platform Performance\n\n*No platform-specific data available.*\n\n"

    section = "## Platform Performance\n\n"
    section += "| Platform | Posts | Reach | Engagement | Eng. Rate | Follower Change |\n"
    section += "|----------|-------|-------|------------|-----------|----------------|\n"

    for platform, data in platforms.items():
        section += (
            f"| {platform.title()} "
            f"| {data.get('posts', 0)} "
            f"| {data.get('reach', 0):,} "
            f"| {data.get('engagement', 0):,} "
            f"| {data.get('engagement_rate', 0):.1f}% "
            f"| {data.get('follower_change', 0):+d} |\n"
        )
    section += "\n"
    return section


def build_top_bottom_content(analytics: dict) -> str:
    """Build top and bottom performing content section."""
    top = analytics.get("top_posts", [])
    bottom = analytics.get("bottom_posts", [])

    section = "## Content Performance Analysis\n\n"

    section += "### Top Performing Posts\n\n"
    if top:
        section += "| Rank | Platform | Type | Engagement Rate | Reach | Key Factor |\n"
        section += "|------|----------|------|----------------|-------|------------|\n"
        for i, post in enumerate(top[:5], 1):
            section += (
                f"| {i} "
                f"| {post.get('platform', 'N/A')} "
                f"| {post.get('format', 'N/A')} "
                f"| {post.get('engagement_rate', 0):.1f}% "
                f"| {post.get('reach', 0):,} "
                f"| {post.get('key_factor', 'N/A')} |\n"
            )
    else:
        section += "*No data available.*\n"
    section += "\n"

    section += "### Bottom Performing Posts\n\n"
    if bottom:
        section += "| Rank | Platform | Type | Engagement Rate | Reach | Diagnosis |\n"
        section += "|------|----------|------|----------------|-------|----------|\n"
        for i, post in enumerate(bottom[:5], 1):
            section += (
                f"| {i} "
                f"| {post.get('platform', 'N/A')} "
                f"| {post.get('format', 'N/A')} "
                f"| {post.get('engagement_rate', 0):.1f}% "
                f"| {post.get('reach', 0):,} "
                f"| {post.get('diagnosis', 'N/A')} |\n"
            )
    else:
        section += "*No data available.*\n"
    section += "\n"

    return section


def build_roi_section(plan: dict, analytics: dict) -> str:
    """Build ROI and business impact section."""
    roi = analytics.get("roi", {})

    section = "## ROI & Business Impact\n\n"

    if roi:
        section += "| Metric | Value |\n"
        section += "|--------|-------|\n"
        for key, value in roi.items():
            display_key = key.replace("_", " ").title()
            section += f"| {display_key} | {value} |\n"
    else:
        section += (
            "ROI metrics are calculated based on tracked conversions and campaign spend. "
            "Connect analytics tracking (UTM parameters, pixel events) for full ROI measurement.\n"
        )
    section += "\n"
    return section


def build_lessons_learned(analytics: dict) -> str:
    """Build lessons learned section."""
    lessons = analytics.get("lessons_learned", [])

    section = "## Lessons Learned\n\n"
    if lessons:
        for i, lesson in enumerate(lessons, 1):
            section += f"{i}. {lesson}\n"
    else:
        section += (
            "Lessons learned are generated from performance pattern analysis. "
            "Key areas to review:\n\n"
            "1. Which content formats consistently outperformed?\n"
            "2. Which posting times drove highest engagement?\n"
            "3. Which content pillars resonated most with the audience?\n"
            "4. Were there any unexpected trends or viral moments?\n"
            "5. What competitive dynamics influenced performance?\n"
        )
    section += "\n"
    return section


def build_recommendations(analytics: dict) -> str:
    """Build next-cycle recommendations section."""
    recs = analytics.get("recommendations", [])

    section = "## Recommendations for Next Campaign Cycle\n\n"
    if recs:
        for i, rec in enumerate(recs, 1):
            category = rec.get("category", "General")
            action = rec.get("action", "")
            rationale = rec.get("rationale", "")
            section += f"### {i}. {category}\n\n"
            section += f"**Action**: {action}\n\n"
            section += f"**Rationale**: {rationale}\n\n"
    else:
        section += (
            "Recommendations are generated from analytics patterns. Key optimization areas:\n\n"
            "- **Content Mix**: Adjust 60/30/10 ratios based on what performed\n"
            "- **Posting Schedule**: Refine day/time based on engagement data\n"
            "- **Format Focus**: Invest more in top-performing formats\n"
            "- **Platform Priority**: Reallocate effort to highest-ROI platforms\n"
            "- **Audience Targeting**: Refine based on engagement demographics\n"
        )
    section += "\n"
    return section


def build_skill_invocation_log(plan: dict) -> str:
    """Build the skill invocation log for pipeline transparency."""
    invocations = plan.get("skill_invocations", [])

    section = "## Pipeline Execution Log\n\n"
    section += "| Phase | Skill | Output | Status |\n"
    section += "|-------|-------|--------|--------|\n"
    for inv in invocations:
        section += (
            f"| Phase {inv['phase']}: {inv['phase_name']} "
            f"| {inv['skill']} "
            f"| {inv['output']} "
            f"| Complete |\n"
        )
    section += "\n"
    return section


# ---------------------------------------------------------------------------
# Report Assembly
# ---------------------------------------------------------------------------

def generate_report(plan: dict, analytics: dict) -> str:
    """Assemble the full campaign report."""
    sections = [
        build_header(plan, analytics),
        build_executive_summary(plan, analytics),
        build_content_overview(plan, analytics),
        build_phase_review(plan, analytics),
        build_platform_breakdown(analytics),
        build_top_bottom_content(analytics),
        build_roi_section(plan, analytics),
        build_lessons_learned(analytics),
        build_recommendations(analytics),
        build_skill_invocation_log(plan),
    ]
    return "\n".join(sections)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(
        description="Generate a campaign performance report from plan and analytics data."
    )
    parser.add_argument(
        "--campaign-plan",
        required=True,
        help="Path to campaign_plan.json",
    )
    parser.add_argument(
        "--analytics",
        required=True,
        help="Path to analytics report JSON (e.g., report_30d.json)",
    )
    parser.add_argument(
        "--output",
        required=True,
        help="Output path for the report",
    )
    parser.add_argument(
        "--format",
        choices=["md", "docx"],
        default="md",
        help="Output format: md (Markdown) or docx (Word). Default: md",
    )
    args = parser.parse_args()

    # Load inputs
    print(f"Loading campaign plan from: {args.campaign_plan}")
    if not os.path.exists(args.campaign_plan):
        print(f"Error: Campaign plan not found: {args.campaign_plan}", file=sys.stderr)
        sys.exit(1)
    with open(args.campaign_plan) as f:
        plan = json.load(f)

    print(f"Loading analytics from: {args.analytics}")
    if not os.path.exists(args.analytics):
        print(f"Warning: Analytics file not found: {args.analytics}", file=sys.stderr)
        print("Generating report with available plan data only.", file=sys.stderr)
        analytics = {}
    else:
        with open(args.analytics) as f:
            analytics = json.load(f)

    # Generate report
    print("Generating campaign report...")
    report_md = generate_report(plan, analytics)

    if args.format == "md":
        # Write Markdown
        os.makedirs(os.path.dirname(os.path.abspath(args.output)), exist_ok=True)
        with open(args.output, "w") as f:
            f.write(report_md)
        print(f"Report written to: {args.output}")

    elif args.format == "docx":
        # Write DOCX using python-docx if available
        try:
            from docx import Document
            from docx.shared import Pt, Inches

            doc = Document()
            style = doc.styles["Normal"]
            style.font.size = Pt(11)

            # Parse markdown into docx (simplified conversion)
            for line in report_md.split("\n"):
                line = line.rstrip()
                if line.startswith("# "):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith("## "):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith("### "):
                    doc.add_heading(line[4:], level=3)
                elif line.startswith("|"):
                    # Tables — add as plain text for simplicity
                    doc.add_paragraph(line, style="Normal")
                elif line.startswith("---"):
                    doc.add_paragraph("_" * 40, style="Normal")
                elif line.startswith("- "):
                    doc.add_paragraph(line[2:], style="List Bullet")
                elif line.startswith("**") and line.endswith("**"):
                    p = doc.add_paragraph()
                    run = p.add_run(line.strip("*"))
                    run.bold = True
                elif line:
                    doc.add_paragraph(line, style="Normal")

            os.makedirs(os.path.dirname(os.path.abspath(args.output)), exist_ok=True)
            doc.save(args.output)
            print(f"Report written to: {args.output}")

        except ImportError:
            print(
                "Warning: python-docx not installed. Falling back to Markdown output.",
                file=sys.stderr,
            )
            fallback_path = args.output.replace(".docx", ".md")
            with open(fallback_path, "w") as f:
                f.write(report_md)
            print(f"Report written to: {fallback_path}")
            print("Install python-docx for DOCX support: pip install python-docx", file=sys.stderr)

    # Summary
    campaign_id = plan.get("campaign_id", "unknown")
    print(f"\nReport generated for campaign: {campaign_id}")
    print(f"Format: {args.format.upper()}")


if __name__ == "__main__":
    main()
